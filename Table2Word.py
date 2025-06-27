import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Length
from openpyxl import load_workbook



def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def excel_sheet_to_word(sheet, word_file):
    doc = Document()
    max_row = sheet.max_row
    max_column = sheet.max_column
    word_table = doc.add_table(rows=max_row, cols=max_column)

    single_spacing = Length(12)  # Word中的单倍行距是12磅

    for excel_row, word_row in zip(sheet.iter_rows(), word_table.rows):
        for excel_cell, word_cell in zip(excel_row, word_row.cells):
            paragraph = word_cell.paragraphs[0]
            run = paragraph.add_run(str(excel_cell.value) if excel_cell.value is not None else '')
            run.font.size = Pt(10.5)
            if isinstance(excel_cell.value, str) and excel_cell.value.isascii():
                run.font.name = 'Times New Roman'
            else:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = single_spacing

    for row_idx, row in enumerate(word_table.rows):
        for col_idx, cell in enumerate(row.cells):
            if row_idx == 0:
                set_cell_border(cell, top={"sz": 12, "color": "000000", "val": "single"})
            if row_idx == 1:
                set_cell_border(cell, top={"sz": 8, "color": "000000", "val": "single"})
            if row_idx == max_row - 1:
                set_cell_border(cell, bottom={"sz": 12, "color": "000000", "val": "single"})
    doc.save(word_file)


if __name__ == "__main__":
    print('Excel to Word转换开始...')
    current_directory = os.getcwd()
    for filename in os.listdir(current_directory):
        if filename.endswith('.xlsx'):
            excel_file = os.path.join(current_directory, filename)
            workbook = load_workbook(excel_file)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                word_file = os.path.splitext(excel_file)[0] + f'_{sheet_name}.docx'
                excel_sheet_to_word(sheet, word_file)
                print(f"Converted {sheet_name} in {filename} to {os.path.basename(word_file)}")
    print('Excel to Word转换完成！')

    os.system('pause')
